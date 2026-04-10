import os
import re
import json
from pathlib import Path
from typing import List, Tuple

import numpy as np
from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    SparseVectorParams,
    SparseIndexParams,
    PointStruct,
    SparseVector,
)

oai = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
qdrant = QdrantClient(host=os.environ.get("QDRANT_HOST", "localhost"), port=6333)

DENSE_DIM = 1536          # text-embedding-3-small
COLLECTION_PREFIX = "doc_"
EMBED_BATCH = 32


# Text extraction 

def extract_text(path: str, ext: str) -> Tuple[str, int]:
    if ext == ".txt":
        text = Path(path).read_text(errors="replace")
        return text, 1

    if ext == ".pdf":
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                pages.append(t)
        return "\n\n".join(pages), len(pages)

    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # also grab table cells
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                if cell_texts:
                    paragraphs.append(" | ".join(cell_texts))
        return "\n\n".join(paragraphs), 1

    raise ValueError(f"Unsupported extension: {ext}")


# Sentence splitting 

def split_sentences(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", text).strip()
    # Split on sentence by whitespace + capital
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z\"\'\(])", text)
    return [s.strip() for s in sentences if len(s.strip()) > 20]


# Semantic chunking 

def embed_sentences(sentences: List[str]) -> np.ndarray:
    all_vecs = []
    for i in range(0, len(sentences), EMBED_BATCH):
        batch = sentences[i : i + EMBED_BATCH]
        resp = oai.embeddings.create(model="text-embedding-3-small", input=batch)
        vecs = [d.embedding for d in resp.data]
        all_vecs.extend(vecs)
    return np.array(all_vecs, dtype=np.float32)


def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-9))


def semantic_chunk(sentences: List[str], vecs: np.ndarray,
                   window: int = 3, threshold_percentile: float = 25) -> List[str]:
    if len(sentences) <= window * 2:
        return [" ".join(sentences)]

    scores = []
    for i in range(window, len(sentences) - window):
        left = vecs[i - window : i].mean(axis=0)
        right = vecs[i : i + window].mean(axis=0)
        scores.append(cosine_sim(left, right))

    cutoff = float(np.percentile(scores, threshold_percentile))
    boundaries = {0}
    for i, s in enumerate(scores):
        if s < cutoff:
            boundaries.add(i + window)
    boundaries.add(len(sentences))

    sorted_bounds = sorted(boundaries)
    chunks = []
    for start, end in zip(sorted_bounds, sorted_bounds[1:]):
        chunk_text = " ".join(sentences[start:end]).strip()
        if chunk_text:
            chunks.append(chunk_text)
    return chunks


# BM25 sparse vectors 

def build_vocab(chunks: List[str]):
    vocab = {}
    for chunk in chunks:
        for token in tokenise(chunk):
            if token not in vocab:
                vocab[token] = len(vocab)
    return vocab


def tokenise(text: str) -> List[str]:
    return re.findall(r"[a-z0-9]+", text.lower())


def bm25_vector(text: str, vocab: dict, corpus_df: dict, N: int,
                k1: float = 1.5, b: float = 0.75, avg_dl: float = 100) -> SparseVector:
    tokens = tokenise(text)
    dl = len(tokens)
    tf: dict[int, float] = {}
    for token in tokens:
        if token in vocab:
            idx = vocab[token]
            tf[idx] = tf.get(idx, 0) + 1

    indices, values = [], []
    for idx, freq in tf.items():
        token = [t for t, i in vocab.items() if i == idx][0]
        df = corpus_df.get(token, 1)
        idf = max(0, np.log((N - df + 0.5) / (df + 0.5) + 1))
        score = idf * (freq * (k1 + 1)) / (freq + k1 * (1 - b + b * dl / avg_dl))
        if score > 0:
            indices.append(idx)
            values.append(float(score))

    return SparseVector(indices=indices, values=values)


# Qdrant helpers 

def ensure_collection(collection: str):
    existing = [c.name for c in qdrant.get_collections().collections]
    if collection not in existing:
        qdrant.create_collection(
            collection_name=collection,
            vectors_config={"dense": VectorParams(size=DENSE_DIM, distance=Distance.COSINE)},
            sparse_vectors_config={
                "sparse": SparseVectorParams(index=SparseIndexParams(on_disk=False))
            },
        )


# Main entry 

def ingest_document(doc_id: str, path: str, ext: str) -> dict:
    text, page_count = extract_text(path, ext)

    sentences = split_sentences(text)
    if not sentences:
        raise ValueError("No extractable text found in document")

    print(f"[ingest] {len(sentences)} sentences extracted")

    vecs = embed_sentences(sentences)
    chunks = semantic_chunk(sentences, vecs)
    print(f"[ingest] {len(chunks)} semantic chunks created")

    # Build BM25 structures
    vocab = build_vocab(chunks)
    corpus_df: dict[str, int] = {}
    for chunk in chunks:
        for token in set(tokenise(chunk)):
            corpus_df[token] = corpus_df.get(token, 0) + 1
    avg_dl = float(np.mean([len(tokenise(c)) for c in chunks]))

    # Embed chunks (dense)
    chunk_vecs = []
    for i in range(0, len(chunks), EMBED_BATCH):
        batch = chunks[i : i + EMBED_BATCH]
        resp = oai.embeddings.create(model="text-embedding-3-small", input=batch)
        chunk_vecs.extend([d.embedding for d in resp.data])

    # Upsert into Qdrant
    collection = COLLECTION_PREFIX + doc_id
    ensure_collection(collection)

    points = []
    for i, (chunk, dvec) in enumerate(zip(chunks, chunk_vecs)):
        svec = bm25_vector(chunk, vocab, corpus_df, len(chunks), avg_dl=avg_dl)
        points.append(
            PointStruct(
                id=i,
                vector={"dense": dvec, "sparse": svec},
                payload={"text": chunk, "chunk_index": i, "doc_id": doc_id},
            )
        )

    qdrant.upsert(collection_name=collection, points=points)
    print(f"[ingest] Upserted {len(points)} points into collection '{collection}'")

    return {
        "chunk_count": len(chunks),
        "page_count": page_count,
        "vocab_size": len(vocab),
        "collection": collection,
    }
